# -*- coding: utf-8 -*-
"""Modellannahmen.docx — kuratierte, verständliche Abgabefassung.

Nur das AKTUELLE Modell (2-Faktor + CET1-Walk-Forward), granular aber
übersichtlich, mit vier didaktischen Schaubildern. Quelle der Fachinhalte:
docs/MODEL_ASSUMPTIONS.md (V2.1) + live SENSITIVITY_MATRIX.
"""
import sys, os, datetime, math
from statistics import NormalDist
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

import pathlib
sys.path.insert(0, str(pathlib.Path(__file__).resolve().parents[2] / "backend"))
from two_factor_stress import SENSITIVITY_MATRIX, stress_pd, stress_lgd

_NORMAL = NormalDist()


def irb_capital_requirement(pd_value, lgd, exposure_class, *, maturity_years=2.5):
    """Scalar Basel-IRB helper for the document's illustrative example."""
    pd_value = min(max(float(pd_value), 1e-12), 1.0 - 1e-12)
    lgd = min(max(float(lgd), 1e-4), 1.0)
    if exposure_class in ("corporate", "bank", "sovereign"):
        r = (1.0 - math.exp(-50.0 * pd_value)) / (1.0 - math.exp(-50.0))
        rho = 0.12 * r + 0.24 * (1.0 - r)
    elif exposure_class == "sme_corporate":
        # KMU-Groessenanpassung mit S = 20 Mio. EUR Jahresumsatz —
        # identisch zu backend/vasicek.rho_sme(pd, sales_m_eur=20.0).
        r = (1.0 - math.exp(-50.0 * pd_value)) / (1.0 - math.exp(-50.0))
        rho = 0.12 * r + 0.24 * (1.0 - r) - 0.04 * (1.0 - (20.0 - 5.0) / 45.0)
    elif exposure_class == "mortgage":
        rho = 0.15
    elif exposure_class == "qrre":
        rho = 0.04
    else:
        r = (1.0 - math.exp(-35.0 * pd_value)) / (1.0 - math.exp(-35.0))
        rho = 0.03 * r + 0.16 * (1.0 - r)
    rho = min(max(rho, 1e-4), 0.999)
    z_alpha = _NORMAL.inv_cdf(0.999)
    cond_pd = _NORMAL.cdf(
        (_NORMAL.inv_cdf(pd_value) + math.sqrt(rho) * z_alpha)
        / math.sqrt(1.0 - rho)
    )
    el = pd_value * lgd
    if exposure_class in ("corporate", "sme_corporate", "bank", "sovereign"):
        m = min(max(float(maturity_years), 1.0), 5.0)
        b = (0.11852 - 0.05478 * math.log(pd_value)) ** 2
        ma = (1.0 + (m - 2.5) * b) / (1.0 - 1.5 * b)
    else:
        ma = 1.0
    k = max((lgd * cond_pd - el) * ma, 0.0)
    return {"K": k, "rwa_density": k * 12.5}

FIGS = os.path.dirname(os.path.abspath(__file__)) + "/figs"
OUT = str(pathlib.Path(__file__).resolve().parents[3] / "Abgabe-Files" / "Abgabedokumente" / "Modellannahmen.docx")


# ---------------------------------------------------------------------
# Live-Kennzahlen aus der Engine. Alle Backtest-Zahlen in Abschnitt 7
# werden beim Bau direkt aus backend/backtesting_walkforward gerechnet;
# nur wenn die lokale Datenbasis fehlt, greifen die dokumentierten
# Fallback-Werte (verifizierter Stand 07.07.2026).
# ---------------------------------------------------------------------
_FALLBACK_STATS = {
    "gesamt": {"n": 30, "mae": 1.26, "bias": -0.99, "naive": 0.70,
               "within": 0.60, "cons": 0.733},
    "jahre": {
        2022: {"n": 10, "brent": 0.099, "rate": 2.77, "mae": 2.49,
               "bias": -2.49, "cons": 1.00, "naive": 0.97},
        2023: {"n": 10, "brent": -0.109, "rate": -0.47, "mae": 0.48,
               "bias": 0.21, "cons": 0.30, "naive": 0.50},
        2024: {"n": 10, "brent": -0.032, "rate": 0.37, "mae": 0.81,
               "bias": -0.70, "cons": 0.90, "naive": 0.64},
    },
    "pd": {"n": 185, "hit": 0.432, "corr": -0.022,
           "mae_model": 0.88, "mae_naive": 0.75},
    "sens": {0.5: {"mae": 0.80, "bias": -0.63, "bias22": -1.31, "cons22": 1.00},
             1.5: {"mae": 1.71, "bias": -1.31, "bias22": -3.55, "cons22": 1.00}},
    "live": False,
}


def _live_stats():
    import pandas as pd_mod
    try:
        import scipy.stats                                      # noqa: F401
    except Exception:                                           # noqa: BLE001
        import types
        import numpy as np

        nd = NormalDist()

        class _Norm:
            @staticmethod
            def cdf(x):
                return np.vectorize(nd.cdf, otypes=[float])(x)

            @staticmethod
            def ppf(x):
                return np.vectorize(nd.inv_cdf, otypes=[float])(x)

            @staticmethod
            def pdf(x):
                return np.vectorize(nd.pdf, otypes=[float])(x)

        scipy_mod = types.ModuleType("scipy")
        stats_mod = types.ModuleType("scipy.stats")
        stats_mod.norm = _Norm()
        scipy_mod.stats = stats_mod
        sys.modules.setdefault("scipy", scipy_mod)
        sys.modules.setdefault("scipy.stats", stats_mod)

    _root = pathlib.Path(__file__).resolve().parents[2]
    for _p in (str(_root), str(_root / "backend")):
        if _p not in sys.path:
            sys.path.insert(0, _p)
    from config import CACHE_DIR, EBA_RAW_DIR                     # type: ignore
    from backtesting_walkforward import (                         # type: ignore
        load_backtest_series, build_cet1_backtest, cet1_backtest_stats,
        compute_annual_macro, build_pd_backtest, pd_backtest_stats)
    from eba_loader import (                                      # type: ignore
        load_historical_capital_panel, panel_to_wide)
    brent = pd_mod.read_parquet(CACHE_DIR / "brent_crude.parquet")
    sven = pd_mod.read_parquet(CACHE_DIR / "svensson_params.parquet")
    cap = panel_to_wide(load_historical_capital_panel(EBA_RAW_DIR))
    ser = load_backtest_series()
    am = compute_annual_macro(brent, sven, [2022, 2023, 2024])

    bt = build_cet1_backtest(cap, ser, am)
    s = cet1_backtest_stats(bt)
    out = {"gesamt": {"n": int(s["n"]), "mae": s["mae_pp"],
                      "bias": s["bias_pp"], "naive": s["mae_naive_pp"],
                      "within": s["within_1pp"],
                      "cons": s["conservative_share"]},
           "jahre": {}, "sens": {}, "live": True}
    for y, g in bt.groupby("year"):
        out["jahre"][int(y)] = {
            "n": len(g),
            "brent": am[int(y)]["d_brent_log"],
            "rate": am[int(y)]["d_r_10y_pp"],
            "mae": float(g["abs_error_pp"].mean()),
            "bias": float(g["error_pp"].mean()),
            "cons": float(g["conservative"].mean()),
            "naive": float((g["cet1_ratio_start"]
                            - g["cet1_ratio_real"]).abs().mean())}
    p = pd_backtest_stats(build_pd_backtest(ser, am))
    out["pd"] = {"n": int(p["n"]), "hit": p["hit_rate"], "corr": p["corr"],
                 "mae_model": p["mae_model"], "mae_naive": p["mae_naive"]}
    # beta-Sensitivitaet: beta*k ist wegen Bilinearitaet der Transmission
    # (dPD = beta*Schock) exakt identisch zu Schock*k.
    for k in (0.5, 1.5):
        am_k = {y: {"d_brent_log": v["d_brent_log"] * k,
                    "d_r_10y_pp": v["d_r_10y_pp"] * k}
                for y, v in am.items()}
        bt_k = build_cet1_backtest(cap, ser, am_k)
        s_k = cet1_backtest_stats(bt_k)
        g22 = bt_k[bt_k["year"] == 2022]
        out["sens"][k] = {"mae": s_k["mae_pp"], "bias": s_k["bias_pp"],
                          "bias22": float(g22["error_pp"].mean()),
                          "cons22": float(g22["conservative"].mean())}
    return out


try:
    STATS = _live_stats()
except Exception as _e:                                            # noqa: BLE001
    print(f"[WARN] Live-Stats nicht verfuegbar ({_e}) — Fallback 07.07.2026")
    STATS = _FALLBACK_STATS


def de(v, d=2, sign=False):
    """Deutsche Zahlformatierung (Komma-Dezimale), optional mit Vorzeichen."""
    s = f"{v:+.{d}f}" if sign else f"{v:.{d}f}"
    return s.replace(".", ",")

NAVY = RGBColor(0x05, 0x1C, 0x2C)
MID = RGBColor(0x03, 0x4B, 0x6F)
GREY = RGBColor(0x6E, 0x6E, 0x6E)

doc = Document()
st = doc.styles["Normal"]
st.font.name = "Arial"; st.font.size = Pt(10.5)
st.element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
st.paragraph_format.space_after = Pt(7)

for hid, size, col in [("Heading 1", 16, NAVY), ("Heading 2", 12.5, MID),
                        ("Heading 3", 11, NAVY)]:
    h = doc.styles[hid]
    h.font.name = "Arial"; h.font.size = Pt(size); h.font.bold = True
    h.font.color.rgb = col
    h.element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")

sec = doc.sections[0]
sec.page_width, sec.page_height = Cm(21.0), Cm(29.7)
for m in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
    setattr(sec, m, Cm(2.2))

foot = sec.footer.paragraphs[0]
foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = foot.add_run("Modellannahmen · EU Banking Credit Stress Cockpit · Seite ")
r.font.size = Pt(8); r.font.color.rgb = GREY; r.font.name = "Arial"
fld = OxmlElement("w:fldSimple"); fld.set(qn("w:instr"), "PAGE")
foot._p.append(fld)


# ---------------- Helpers ----------------
def shade(el, hexcol):
    sh = OxmlElement("w:shd"); sh.set(qn("w:val"), "clear")
    sh.set(qn("w:fill"), hexcol); el.append(sh)


def body(text, bold_parts=(), size=10.5, italic=False, color=None,
         space_after=7):
    """Absatz; **…** wird fett gesetzt."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    for i, tok in enumerate(text.split("**")):
        if not tok:
            continue
        r = p.add_run(tok)
        r.bold = (i % 2 == 1)
        r.italic = italic
        r.font.size = Pt(size)
        if color:
            r.font.color.rgb = color
    return p


def bullet(text):
    p = doc.add_paragraph(style="List Bullet")
    for i, tok in enumerate(text.split("**")):
        if not tok:
            continue
        r = p.add_run(tok); r.bold = (i % 2 == 1); r.font.size = Pt(10.5)
    return p


def h1(t): doc.add_paragraph(t, style="Heading 1")
def h2(t): doc.add_paragraph(t, style="Heading 2")


def figure(fname, width_cm, caption):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(2)
    p.add_run().add_picture(f"{FIGS}/{fname}", width=Cm(width_cm))
    c = doc.add_paragraph(); c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    c.paragraph_format.space_after = Pt(12)
    r = c.add_run(caption); r.font.size = Pt(8.5); r.font.color.rgb = GREY
    r.italic = True


def table(headers, rows, col_pct, fontsize=9):
    tab = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tab.style = "Table Grid"; tab.alignment = WD_TABLE_ALIGNMENT.CENTER
    total = Cm(16.6)
    for c, htxt in enumerate(headers):
        cell = tab.rows[0].cells[c]
        cell.width = Cm(16.6 * col_pct[c] / 100)
        cell.paragraphs[0].text = ""
        r = cell.paragraphs[0].add_run(htxt)
        r.bold = True; r.font.size = Pt(fontsize)
        shade(cell._tc.get_or_add_tcPr(), "D9E4EC")
    for ri, row in enumerate(rows):
        for c, val in enumerate(row):
            cell = tab.rows[1 + ri].cells[c]
            cell.width = Cm(16.6 * col_pct[c] / 100)
            cell.paragraphs[0].text = ""
            for i, tok in enumerate(str(val).split("**")):
                if not tok:
                    continue
                r = cell.paragraphs[0].add_run(tok)
                r.bold = (i % 2 == 1); r.font.size = Pt(fontsize)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return tab


# ---------------- Titelseite ----------------
t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
t.paragraph_format.space_before = Pt(150)
r = t.add_run("Modellannahmen"); r.font.size = Pt(32); r.bold = True
r.font.color.rgb = NAVY
s2 = doc.add_paragraph(); s2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = s2.add_run("EU Banking Credit Stress Cockpit\n"
               "2-Faktor-Kreditstress (Basel-III-IRB) · Zielgröße CET1-Quote · "
               "Pillar-3-Datenbasis")
r.font.size = Pt(13); r.font.color.rgb = MID
meta = doc.add_paragraph(); meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.paragraph_format.space_before = Pt(30)
r = meta.add_run("Kuratierte Abgabefassung · Stand "
                 + datetime.date.today().strftime("%d.%m.%Y")
                 + "\nTechnische Referenz: Modellarchitektur/docs/MODEL_ASSUMPTIONS.md (V2.1)")
r.font.size = Pt(9.5); r.font.color.rgb = GREY
doc.add_page_break()

# ---------------- Inhaltsverzeichnis ----------------
p = doc.add_paragraph(); r = p.add_run("Inhaltsverzeichnis")
r.font.size = Pt(16); r.bold = True; r.font.color.rgb = NAVY
table(
    ["Kapitel", "Seite"],
    [["Abkürzungen und Symbole", "3"],
     ["1 · Das Modell auf einen Blick", "5"],
     ["  Prüfpfad für kritische Evaluation", "5"],
     ["2 · Die zwei Risikofaktoren", "6"],
     ["  Warum Brent-log-Return?", "6"],
     ["  Datenquellen und ihre Rollen", "7"],
     ["3 · Sektor-differenzierte Stress-Transmission", "7"],
     ["  Ökonomische Logik je Kreditklasse", "8"],
     ["  Durchgerechnetes Beispiel", "8"],
     ["4 · Kapitalrechnung und CET1-Bridge", "9"],
     ["  Die Formelstrecke im Detail", "9"],
     ["  Der Sovereign-Kanal konkret", "10"],
     ["5 · Datenbasis und Integritätssicherung", "11"],
     ["6 · Annahmen-Inventar (A-01 bis A-10)", "12"],
     ["7 · Validierung: CET1-Walk-Forward-Backtest", "15"],
     ["8 · Bewusste Scope-Grenzen", "16"],
     ["9 · Bibliographie", "17"]],
    [86, 14], fontsize=9)
doc.add_page_break()

# ---------------- Abkürzungen und Symbole ----------------
h1("Abkürzungen und Symbole")
body("Kurzreferenz für die kritische Durchsicht — jede Abkürzung wird "
     "zusätzlich bei ihrem ersten inhaltlichen Auftreten im Text erläutert.")
h2("Abkürzungen")
table(
    ["Abkürzung", "Bedeutung"],
    [["A-IRB / F-IRB", "Advanced / Foundation Internal Ratings-Based — im "
      "A-IRB schätzt die Bank PD und LGD selbst, im F-IRB nur die PD "
      "(LGD regulatorisch, z. B. 45 % für unbesicherte Senior-Forderungen)"],
     ["AC", "Amortised Cost — zu fortgeführten Anschaffungskosten "
      "bilanziert; kein laufender Marktwert-Effekt im Eigenkapital"],
     ["ASRF", "Asymptotic Single Risk Factor — Portfoliomodell hinter der "
      "Basel-IRB-Formel (Vasicek/Gordy)"],
     ["bp", "Basispunkt = 0,01 Prozentpunkte"],
     ["CCF", "Credit Conversion Factor — Umrechnungsfaktor für "
      "außerbilanzielle Kreditzusagen in EAD"],
     ["CET1", "Common Equity Tier 1 — hartes Kernkapital; die CET1-Quote "
      "(CET1 ÷ RWA) ist die Zielgröße des Modells"],
     ["CRR", "Capital Requirements Regulation — Verordnung (EU) Nr. 575/2013"],
     ["EAD", "Exposure at Default — Forderungshöhe bei Ausfall (post-CCF)"],
     ["EBA", "European Banking Authority — EU-Bankenaufsichtsbehörde"],
     ["EL", "Expected Loss — erwarteter Verlust = PD · LGD · EAD"],
     ["EU CR6", "Pillar-3-Offenlegungstemplate für IRB-Kreditrisiko "
      "(PD/LGD/EAD je Exposure-Klasse; EBA ITS/2020/04)"],
     ["FVOCI", "Fair Value through Other Comprehensive Income — "
      "Marktwertänderungen wirken erfolgsneutral direkt im Eigenkapital (OCI)"],
     ["FVTPL / HfT", "Fair Value through Profit or Loss / Held for Trading — "
      "Marktwertänderungen wirken über die Gewinn- und Verlustrechnung"],
     ["ICAAP", "Internal Capital Adequacy Assessment Process — "
      "bankinterner Kapitaladäquanz-Prozess"],
     ["IRB", "Internal Ratings-Based Approach — auf internen Ratings "
      "basierender Ansatz der Kapitalunterlegung"],
     ["LEI", "Legal Entity Identifier — eindeutige Kennung je Institut"],
     ["LGD", "Loss Given Default — Verlustquote bei Ausfall"],
     ["MAE", "Mean Absolute Error — mittlerer absoluter Prognosefehler"],
     ["MtM", "Mark-to-Market — Bewertung zum aktuellen Marktwert"],
     ["NII / NIM", "Net Interest Income / Margin — Zinsüberschuss bzw. "
      "Zinsmarge einer Bank"],
     ["OCI", "Other Comprehensive Income — erfolgsneutrales Eigenkapital "
      "(sonstiges Ergebnis)"],
     ["PD", "Probability of Default — 1-Jahres-Ausfallwahrscheinlichkeit "
      "(CRR Art. 180)"],
     ["PIT / TTC", "Point-in-Time / Through-the-Cycle — stichtagsbezogene "
      "vs. über den Zyklus geglättete Risikomessung"],
     ["pp", "Prozentpunkte"],
     ["QRRE", "Qualifying Revolving Retail Exposures — revolvierende "
      "Konsumentenkredite (z. B. Kreditkarten)"],
     ["RWA", "Risk-Weighted Assets — risikogewichtete Aktiva; Nenner der "
      "CET1-Quote"],
     ["SREP", "Supervisory Review and Evaluation Process — aufsichtlicher "
      "Überprüfungsprozess"]],
    [17, 83], fontsize=8.5)
h2("Symbole")
table(
    ["Symbol", "Bedeutung", "Ort"],
    [["ΔBrent_log", "Ölpreis-Schock als log-Return ln(P₁/P₀); "
      "+0,50 ≈ +65 % Preis", "Abschnitt 2"],
     ["Δr_10y", "Parallelverschiebung des 10-Jahres-Zinses in pp",
      "Abschnitt 2"],
     ["β_Öl, β_Zins", "PD-Sensitivitäten je Kreditklasse "
      "(pp ΔPD pro Faktoreinheit)", "Abschnitt 3"],
     ["γ_Öl, γ_Zins", "LGD-Sensitivitäten je Kreditklasse", "Abschnitt 3"],
     ["ρ", "Asset-Korrelation der IRB-Formel (CRR Art. 153/154)",
      "Abschnitt 4"],
     ["K", "Kapitalanforderung je Einheit Exposure (IRB-Formel)",
      "Abschnitt 4"],
     ["N, N⁻¹", "Standardnormal-Verteilungsfunktion und ihre Inverse",
      "Abschnitt 4"],
     ["α = 99,9 %", "Konfidenzniveau der IRB-Formel", "Abschnitt 4"],
     ["b(PD), MA", "Laufzeit-Glättungsfunktion und Laufzeit-Adjustment",
      "Abschnitt 4"],
     ["M", "effektive Restlaufzeit; Konvention M = 2,5 Jahre",
      "Abschnitt 4"],
     ["S", "Jahresumsatz für die KMU-ρ-Anpassung; Konvention 20 Mio. €",
      "Abschnitt 4"],
     ["D_bucket", "Modified Duration je Laufzeit-Bucket (Sovereign-Kanal)",
      "Abschnitt 4"],
     ["ΔRWA, ΔEL, ΔMtM", "Stress-Deltas der CET1-Bridge", "Abschnitt 4"],
     ["e_i, MAE, Bias", "Prognosefehler je Bank-Jahr und Fehlermaße",
      "Abschnitt 7"]],
    [16, 62, 22], fontsize=8.5)
doc.add_page_break()

# ================= 1 · Das Modell auf einen Blick =================
h1("1 · Das Modell auf einen Blick")
body("Das Cockpit beantwortet eine aufsichtlich zentrale Frage: "
     "**Wie stünde die harte Kernkapitalquote (CET1) der zehn größten "
     "EU-IRB-Banken, wenn ein Ölpreis- und/oder Zinsschock einträte?** "
     "Die CET1-Quote ist die Solvenz-Kennzahl der Bankenaufsicht — fällt sie "
     "unter die gestaffelten Mindestanforderungen, greifen aufsichtliche "
     "Maßnahmen bis hin zu Ausschüttungssperren.")
body("Vier Prinzipien tragen das Modell:")
bullet("**Quellenreine Datenbasis** — alle Risikoparameter (PD, LGD, EAD) "
       "stammen aus den offiziellen Pillar-3-Offenlegungen (EU CR6) der "
       "Banken selbst; kein Wert wird geschätzt oder abgeleitet.")
bullet("**Zwei getrennte, messbare Faktoren** — Ölpreis (ΔBrent) und "
       "10-Jahres-Zins (Δr_10y) werden nicht zu einem Kunstfaktor "
       "aggregiert, sondern wirken einzeln und beobachtbar.")
bullet("**Sektor-differenzierte Transmission** — jede Kreditklasse reagiert "
       "gemäß quellenbelegter Sensitivitäten unterschiedlich (eine "
       "Hypothek anders als ein Kreditkartenbuch).")
bullet("**Konservative Auslegung** — im Zweifel überschätzt das Modell den "
       "Kapitalbedarf; entlastende Effekte (z. B. Zinsüberschuss) werden "
       "bewusst nicht gegengerechnet.")
figure("fig1_pipeline.png", 16.2,
       "Abb. 1 · Modell-Pipeline: Zwei Makro-Schocks laufen über die "
       "β-Transmission und die Basel-III-Kapitalformel in das Kreditbuch; "
       "der Zinsschock wirkt zusätzlich über den Sovereign-Kanal. Zielgröße "
       "ist stets die CET1-Quote unter Stress.")
h2("Prüfpfad für kritische Evaluation")
body("Das Dokument ist so aufgebaut, dass ein Dritter die Modelllogik vom "
     "Rohdatum bis zur CET1-Quote nachvollziehen kann. Die folgende "
     "Kurzlandkarte nennt die jeweils prüfbare Stelle; Details folgen in "
     "den Abschnitten 2 bis 8.")
table(
    ["Prüffrage", "Nachvollziehbarer Beleg", "Abschnitt"],
    [["Woher kommen PD, LGD und EAD?",
      "Pillar-3-EU-CR6-Zeilen mit Quellen-URL, Seitenzahl und Tabellenname; "
      "keine Ableitung aus RWA oder NPL-Quoten",
      "5 / A-01 / A-02"],
     ["Welche Werte sind Annahmen, welche nicht?",
      "Konfidenzklassen [published], [estimate], [approximation], "
      "[assumption]; ρ ist regulatorischer CRR-Parameter, keine Schätzung",
      "4 / 6"],
     ["Wie läuft ein Schock durch das Modell?",
      "ΔBrent und Δr_10y → ΔPD/ΔLGD → IRB-K-Formel → ΔRWA/ΔEL → CET1-Bridge",
      "3 / 4"],
     ["Wie wird Look-ahead im Backtest verhindert?",
      "Walk-Forward: Portfolio zum Vorjahresstichtag einfrieren, realen "
      "Jahresschock einspeisen, erst danach mit gemeldeter CET1 vergleichen",
      "7 / A-02c"],
     ["Wo liegen die bewussten Grenzen?",
      "Scope-Grenzen separat ausgewiesen: u. a. kein NII-Gegeneffekt, "
      "kein Hedge-Rekonstrukt, kein Mehrjahrespfad",
      "8"]],
    [30, 50, 20], fontsize=8.5)

# ================= 2 · Die zwei Risikofaktoren =================
h1("2 · Die zwei Risikofaktoren")
body("Beide Faktoren sind täglich beobachtbare Marktgrößen — das Modell "
     "braucht keine latenten Zustandsvariablen:")
table(
    ["Faktor", "Einheit", "Ökonomische Wirkung", "Lesebeispiel"],
    [["**ΔBrent** (Ölpreis)", "log-Return",
      "Energie-/Kostenschock: verteuert Vorleistungen, drückt "
      "Schuldner-Cashflows und Haushaltsbudgets",
      "+0,50 ≈ +65 % Ölpreis"],
     ["**Δr_10y** (10-J-Zins)", "Prozentpunkte",
      "Refinanzierungs-/Bewertungsschock: verteuert "
      "Anschlussfinanzierung, senkt Sicherheiten- und Anleihewerte",
      "+2,0 = +200 Basispunkte"]],
    [16, 12, 50, 22])
h2("Warum Brent-log-Return?")
body("**Warum Brent?** Brent ist der liquide, täglich verfügbare "
     "Rohöl-Benchmark mit hoher Relevanz für europäische Energie- und "
     "Importpreise. Das Modell nutzt Brent nicht, weil jede Bank direkt "
     "Ölpreisrisiko hält, sondern als beobachtbaren Proxy für den "
     "Energie-/Angebots-/Inflationsschock, der Unternehmensmargen, "
     "Haushaltsbudgets und Sicherheitenwerte belastet. WTI wäre stärker "
     "durch US-spezifische Lager- und Transportbedingungen geprägt; Gas- "
     "oder Strompreise wären für Europa ebenfalls relevant, aber deutlich "
     "heterogener und weniger stabil als ein einheitlicher, täglich "
     "verfügbarer Faktor.")
body("**Warum log-Return statt einfacher Prozentänderung?** Der Ölpreis "
     "wird als ΔBrent_log = ln(P_t / P_0) gemessen. Log-Returns sind "
     "skalierungsfrei, über Zeiträume addierbar und behandeln relative "
     "Preisbewegungen konsistent: +0,50 entspricht exp(0,50)−1 ≈ +65 %, "
     "−0,50 entspricht etwa −39 %. Gerade bei großen Stressbewegungen "
     "vermeidet die Log-Konvention Verzerrungen einfacher Prozentänderungen "
     "und ist Standard in der Finanzzeitreihenanalyse. Die β-Werte sind "
     "daher als **Prozentpunkte ΔPD/ΔLGD pro Logpunkt Brent-Schock** zu "
     "lesen. Der Zinsfaktor wird dagegen in Prozentpunkten gemessen, weil "
     "Renditen bereits Prozentgrößen sind.")
body("**Warum zwei getrennte Faktoren?** Gemessen auf Tagesbasis — "
     "tägliche Brent-log-Returns gegen tägliche Änderungen des "
     "10-Jahres-Svensson-Zinses, rollierendes 5-Jahres-Fenster mit "
     "1 242 Handelstagen — beträgt die Pearson-Korrelation der beiden "
     "Faktoren nur ρ = +0,07 (Fisher-z-95 %-Konfidenzintervall "
     "[+0,01; +0,12]; OLS-R² = 0,005). Sie tragen also nahezu unabhängige "
     "Information — eine Aggregation zu einem Einheitsfaktor würde genau "
     "die Unterscheidung wegwerfen, die z. B. eine zinsgetriebene "
     "Hypothekenkrise von einem Energiekostenschock trennt. Die Grenze "
     "dieser Tages-Evidenz für den Jahres-Horizont diskutiert Annahme "
     "A-10; die Slider-Spannen des Cockpits (ΔBrent_log ∈ [−1,0; +1,5], "
     "Δr_10y ∈ [−3; +5] pp) orientieren sich an der Größenordnung des "
     "adversen EBA/ESRB-2025-Szenarios (§4.1.6).")
h2("Datenquellen und ihre Rollen")
table(
    ["Quelle", "Inhalt", "Rolle im Modell"],
    [["Pillar-3-Reports (EU CR6)", "PD, LGD, EAD je Bank × Kreditklasse",
      "**Input** — Risikoparameter (Live-Snapshot 31.12.2024 + "
      "Backtest-Reihe FY2021–2024)"],
     ["EBA Transparency Exercise", "CET1, RWA je Bank und Quartal",
      "**Vergleichsseite** — realisierte Werte für die Validierung; nie "
      "PD/LGD-Quelle"],
     ["Brent (ICE), täglich", "Ölpreis-Historie", "Faktor 1"],
     ["Bundesbank-Svensson-Parameter", "Zinsstrukturkurve, täglich "
      "(Nelson/Siegel 1987, erweitert um Svensson 1994)",
      "Faktor 2 (10-J-Zins)"],
     ["EBA IFRS-9-Split", "HfT/FVTPL/FVOCI-Anteile der Staatsanleihen",
      "Sovereign-Kanal (CET1-wirksamer Anteil)"]],
    [24, 33, 43])

# ================= 3 · Stress-Transmission =================
h1("3 · Sektor-differenzierte Stress-Transmission")
body("Kern des Modells ist die Übersetzung der beiden Schocks in "
     "Risikoparameter — **mit eigenen Sensitivitäten je Kreditklasse** "
     "(Annahme A-04):")
body("ΔPD [pp] = β_Öl · ΔBrent + β_Zins · Δr_10y      "
     "ΔLGD [pp] = γ_Öl · ΔBrent + γ_Zins · Δr_10y",
     size=10.5, color=MID)
body("Die gestressten Werte werden begrenzt — jede Grenze hat einen "
     "Grund: Die **PD-Untergrenze von 0,03 %** ist der regulatorische "
     "Basel-Floor (CRR Art. 160/163 — keine Bank darf einem lebenden "
     "Portfolio eine geringere Ausfallwahrscheinlichkeit zuweisen). Die "
     "**PD-Obergrenze von 50 %** dient der numerischen Stabilität: jenseits "
     "davon wäre ein Portfolio faktisch im Ausfallstatus und die "
     "IRB-Formel nicht mehr die richtige Linse. Die **LGD-Untergrenze von "
     "5 %** verhindert unplausibel verlustfreie Ausfälle selbst bei bester "
     "Besicherung; die **Obergrenze 100 %** ist definitorisch (mehr als "
     "das Exposure kann nicht verloren gehen).")
figure("fig2_beta.png", 14.5,
       "Abb. 2 · Die Sensitivitäts-Matrix als Bild: Jede Zelle ist eine "
       "dokumentierte Annahme — rot = Risiko steigt mit dem Schock, blau = "
       "sinkt, weiß = kein Effekt. Auffällig: Hypotheken sind zinsgetrieben, "
       "Banken profitieren leicht von Zinsanstiegen (Zinsüberschuss), "
       "Staaten reagieren nicht über das Kreditbuch (separater "
       "Sovereign-Kanal).")
h2("Ökonomische Logik je Kreditklasse")
_logik = {
    "corporate": "Energie verteuert Produktion, Zinsen die Refinanzierung; "
                 "LGD steigt mit fallenden Sicherheitenwerten.",
    "sme_corporate": "Wie Corporate, aber ≈2× stärker — KMU haben dünnere "
                     "Puffer und weniger Preissetzungsmacht (ECB WP 2897).",
    "mortgage": "Zins dominiert: Raten steigen (Floating/Anschluss), "
                "Hauspreise fallen → höchste LGD-Zins-Sensitivität.",
    "qrre": "Kreditkarten/Revolving: hängt am Haushaltsbudget → "
            "energie-/inflationssensitiv, kaum zins-sensitiv.",
    "other_retail": "Konsum-/Autokredite: Mischprofil zwischen Mortgage "
                    "und QRRE.",
    "bank": "Leicht **negatives** Zins-β: steigende Zinsen stützen den "
            "Zinsüberschuss der Gegenpartei-Banken (NIM-Effekt).",
    "sovereign": "Im Kreditbuch bewusst inert — der Zins wirkt auf Staaten "
                 "über den Marktwert-Kanal (Abschnitt 4), nicht über die PD.",
}
rows = []
for c in ["corporate", "sme_corporate", "mortgage", "qrre",
          "other_retail", "bank", "sovereign"]:
    m = SENSITIVITY_MATRIX[c]
    fmt = lambda v: f"{v:+.2f}".replace(".", ",")
    rows.append([c, fmt(m['pd_rate']), fmt(m['pd_oil']),
                 fmt(m['lgd_rate']), fmt(m['lgd_oil']), _logik[c]])
table(["Klasse", "β PD·Zins", "β PD·Öl", "γ LGD·Zins", "γ LGD·Öl",
       "Ökonomische Logik (Kurzform)"],
      rows, [14, 9, 9, 10, 9, 49], fontsize=8.5)
body("Quellen der Kalibrierung: EBA (2024) *2025 EU-wide Stress Test — "
     "Methodological Note* §2.4.2 (sektorale Sensitivitäten sind der "
     "aufsichtlich vorgesehene Ansatz); ECB WP 2897 (Unternehmens-PD auf "
     "Angebots-/Zinsschocks, KMU-Verstärkung ≈ 2–3×), ECB WP 3207 "
     "(Sektor-Heterogenität der Unternehmens-Ausfälle im Stress), ECB WP "
     "3112 (Hypotheken-Defaults und Zins), ECB Financial Stability Review "
     "Mai 2024 (Haushalte); die relative Rangfolge der Klassen ist gegen "
     "EBA Results 2025, Fig. 22 (projizierte Verlustquoten je Portfolio) "
     "quervalidiert. Die β sind als belegte, per Regler übersteuerbare "
     "Defaults ausgewiesen — Konfidenzklasse [estimate]; ihre Wirkung auf "
     "das Endergebnis quantifiziert der ±50-%-Robustheitstest in "
     "Abschnitt 7.", size=9.5,
     italic=True)

h2("Durchgerechnetes Beispiel: ein Schock läuft durch das Modell")
# Live aus der Engine gerechnet — keine handgepflegten Zahlen.
_pd0, _lgd0 = 0.025, 0.35
_dr, _db = 2.0, 0.30
_pd1 = stress_pd(_pd0, _db, _dr, "corporate")
_lgd1 = stress_lgd(_lgd0, _db, _dr, "corporate")
_r0 = irb_capital_requirement(_pd0, _lgd0, "corporate", maturity_years=2.5)
_r1 = irb_capital_requirement(_pd1, _lgd1, "corporate", maturity_years=2.5)
_g = lambda v, d=2: f"{v:.{d}f}".replace(".", ",")
_mx = SENSITIVITY_MATRIX["corporate"]
body("Angenommen wird ein Corporate-Segment mit PD = 2,50 % und "
     "LGD = 35,00 %; als Szenario treten ein Zinsanstieg von +2,0 "
     "Prozentpunkten und ein Brent-log-Return von +0,30 (entspricht "
     f"e^0,30 − 1 ≈ +{_g((math.exp(_db)-1)*100, 0)} % Ölpreis, vgl. die "
     "Log-Konvention aus Abschnitt 2) gleichzeitig ein. Dann rechnet das "
     "Modell in drei Schritten:")
bullet(f"**Schritt 1 — PD-Transmission:** ΔPD = {_g(_mx['pd_rate'])} · 2,0 "
       f"+ {_g(_mx['pd_oil'])} · 0,30 = "
       f"**+{_g((_pd1-_pd0)*100)} pp** → gestresste PD "
       f"**{_g(_pd1*100)} %**.")
bullet(f"**Schritt 2 — LGD-Transmission:** ΔLGD = {_g(_mx['lgd_rate'])} · "
       f"2,0 + {_g(_mx['lgd_oil'])} · 0,30 = "
       f"**+{_g((_lgd1-_lgd0)*100)} pp** → gestresste LGD "
       f"**{_g(_lgd1*100)} %**.")
bullet("**Schritt 3 — Kapitalwirkung (IRB-Formel):** der erwartete Verlust "
       f"steigt von {_g(_pd0*_lgd0*100, 3)} % auf {_g(_pd1*_lgd1*100, 3)} % "
       f"des Exposures (relativ **+{_g((_pd1*_lgd1/(_pd0*_lgd0)-1)*100, 0)} %**), "
       "die RWA-Dichte von "
       f"{_g(float(_r0['rwa_density'])*100, 1)} % auf "
       f"{_g(float(_r1['rwa_density'])*100, 1)} % — ein Anstieg um "
       f"**+{_g((float(_r1['rwa_density'])-float(_r0['rwa_density']))*100, 1)} pp** "
       f"absolut (relativ "
       f"+{_g((float(_r1['rwa_density'])/float(_r0['rwa_density'])-1)*100, 1)} %).")
body("Auf Portfolioebene wiederholt sich diese Rechnung für jede Bank und "
     "jede Klasse mit den echten Pillar-3-Parametern; die Summe der "
     "ΔRWA- und ΔEL-Beiträge speist die CET1-Bridge (Abschnitt 4). Die "
     "Beispielwerte sind live aus der Modell-Engine gerechnet, nicht "
     "handgepflegt.", size=9.5, italic=True)

# ================= 4 · Kapitalrechnung =================
h1("4 · Kapitalrechnung und CET1-Bridge")
body("Die gestressten PD/LGD laufen durch die **regulatorische "
     "Basel-III-IRB-Kapitalformel** (Vasicek/ASRF; BCBS 2017, §272–284) — "
     "bewusst der aufsichtliche Standard, keine Eigenkonstruktion:")
body("Öl-/Zinsschock → PD_stress, LGD_stress → Basel-IRB-K-Formel → RWA → CET1",
     size=10, color=MID)
body("K = [ LGD · N( (N⁻¹(PD) + √ρ · N⁻¹(0,999)) / √(1−ρ) ) − PD · LGD ] · MA"
     "      →      RWA = K · 12,5 · EAD", size=10, color=MID)
body("CET1-Quote_stress = ( CET1 − ΔEL_Kredit + ΔMtM_Sov ) / "
     "( RWA_total + ΔRWA_Kredit )", size=10, color=MID)
body("Dabei ist ρ der regulatorische IRB-Klassenparameter nach CRR "
     "Art. 153/154, MA das Laufzeit-Adjustment und N die "
     "Standardnormalverteilung; ΔMtM_Sov ist bei Zinsanstieg negativ "
     "(Kurswertverlust) und umfasst nur den CET1-wirksamen "
     "FVOCI/FVTPL-Anteil. Das Kreditbuch wirkt also auf Zähler (−ΔEL) "
     "**und** Nenner (+ΔRWA), der Sovereign-Kanal nur auf den Zähler. "
     "Die theoretische Fundierung des ASRF-Rahmens liefern Vasicek "
     "(1991, 2002) und Gordy (2003).")

h2("Die Formelstrecke im Detail")
body("Für die kritische Evaluation sind vier Formeln zentral. Sie bilden "
     "die komplette Wirkungskette ab: zuerst werden die makroökonomischen "
     "Schocks in Risikoparameter übersetzt, danach werden diese Parameter "
     "in regulatorische Kapitalanforderungen und schließlich in die CET1-"
     "Quote überführt. Wichtig: Die Formeln trennen strikt zwischen "
     "**gemeldeten Größen** (PD, LGD, EAD, CET1, RWA, IFRS-9-Split), "
     "**regulatorisch vorgegebenen Parametern** (z. B. ρ, 99,9 %-Quantil, "
     "RWA-Faktor 12,5) und **Modellkonventionen** (z. B. lineare β/γ-"
     "Transmission, konstante EAD unter Stress).")
body("**1 · Stress-Transmission:** "
     "PD_stress = clip(PD_0 + β_Zins · Δr_10y + β_Öl · ΔBrent_log); "
     "LGD_stress = clip(LGD_0 + γ_Zins · Δr_10y + γ_Öl · ΔBrent_log). "
     "Diese Gleichungen sind die eigentliche ökonomische Modellannahme: "
     "sie legen fest, wie stark eine Kreditklasse auf Öl- und Zinsschocks "
     "reagiert. Die Grenzen verhindern unplausible oder numerisch instabile "
     "Werte (PD-Floor 0,03 %, PD-Cap 50 %, LGD 5–100 %).")
body("**2 · Erwarteter Verlust:** EL = PD · LGD · EAD; entsprechend "
     "ΔEL = EL_stress − EL_0. Der erwartete Verlust ist der direkte "
     "Zähler-Effekt der Kreditbuch-Bridge: zusätzlicher erwarteter Verlust "
     "senkt das harte Kernkapital. Diese Behandlung ist bewusst einfach und "
     "konservativ; sie ersetzt keine vollständige IFRS-9-Lifetime-EL-"
     "Simulation, sondern bildet den Einjahres-Stress als Risikovorsorge-"
     "Proxy ab.")
body("**3 · Unerwarteter Verlust / RWA:** K ist die IRB-Kapitalanforderung "
     "pro Exposure-Einheit. Der Term mit N und N⁻¹ transformiert die "
     "unbedingte PD in eine 99,9-%-Stress-PD des asymptotischen "
     "Ein-Faktor-Portfolios; LGD multipliziert diese Stress-PD und der "
     "erwartete Verlust wird abgezogen. Aus K folgt RWA = K · 12,5 · EAD, "
     "weil 8 % Mindestkapitalquote dem Multiplikator 1 / 0,08 = 12,5 "
     "entsprechen.")
body("**4 · CET1-Bridge:** CET1_stress = (CET1_0 − ΔEL_Kredit + "
     "ΔMtM_Sovereign) / (RWA_0 + ΔRWA_Kredit). Damit ist klar, welcher "
     "Kanal wo wirkt: Kreditrisiko senkt den Zähler und erhöht den Nenner; "
     "Sovereign-MtM wirkt nur auf den Zähler und nur für Positionen, deren "
     "Bewertung nach IFRS 9 tatsächlich CET1-wirksam ist.")
body("**Laufzeit-Adjustment:** b(PD) = (0,11852 − 0,05478 · ln PD)² und "
     "MA = (1 + (M − 2,5) · b) / (1 − 1,5 · b). MA gilt für corporate, "
     "sme_corporate, bank und sovereign; für die Retail-Klassen "
     "(mortgage, qrre, other_retail) ist MA = 1 (CRR-Systematik). "
     "**Zwei stille Konventionen, hier explizit ausgewiesen "
     "[assumption]:** Die effektive Restlaufzeit ist einheitlich "
     "**M = 2,5 Jahre** gesetzt (Basel-Default; Pillar 3 legt keine "
     "Restlaufzeit je Klasse offen — konstant über alle Jahrgänge und "
     "damit zeitvergleichs-neutral), und die KMU-Größenanpassung nutzt "
     "einen repräsentativen Jahresumsatz **S = 20 Mio. €** für das "
     "aggregierte SME-Buch.")
body("**Kanal-Zerlegung:** Die Stress-Wirkung wird sequentiell zerlegt — "
     "Stufe 1 stresst nur die PD (PD-Kanal), Stufe 2 zusätzlich die LGD "
     "(LGD-Kanal); beide Teilbeträge addieren sich exakt zu ΔRWA bzw. "
     "ΔEL. **ΔEL als Zähler-Proxy:** ΔEL = Δ(PD · LGD · EAD) wird 1:1 "
     "vom CET1-Zähler abgezogen — als Proxy der stressbedingten "
     "Risikovorsorge. Die regulatorische EL-Shortfall-Mechanik (CRR "
     "Art. 36, 62) wird nicht separat modelliert — eine bewusste "
     "Vereinfachung auf der konservativen Seite.")
body("**Asset-Korrelation ρ je Klasse** — regulatorisch fixiert (CRR "
     "Art. 153/154), keine Modellannahme:")
table(
    ["Klasse", "ρ-Festlegung (wie implementiert)", "Spanne / Wert"],
    [["corporate, bank, sovereign",
      "ρ = 0,12·R + 0,24·(1−R) mit R = (1−e^(−50·PD)) / (1−e^(−50))",
      "0,24 → 0,12 (fallend in PD)"],
     ["sme_corporate",
      "wie corporate, abzüglich Größenanpassung 0,04 · (1−(S−5)/45) "
      "bei S = 20 Mio. €",
      "Abschlag 0,027"],
     ["mortgage", "konstant", "0,15"],
     ["qrre", "konstant", "0,04"],
     ["other_retail",
      "ρ = 0,03·R₃₅ + 0,16·(1−R₃₅) mit R₃₅ analog (Faktor 35)",
      "0,16 → 0,03 (fallend in PD)"]],
    [24, 52, 24], fontsize=8.5)

h2("Der Sovereign-Kanal konkret")
body("Parallel zum Kreditbuch werden die Staatsanleihebestände bewertet: "
     "**ΔMtM = −D_bucket · Δr · Exposure_bucket**, summiert über die "
     "EBA-Laufzeit-Buckets je Bank. Die Modified Duration je Bucket ist "
     "eine Konvention über den Bucket-Mittelpunkt — 0,125 / 0,625 / 1,5 / "
     "2,5 / 4,0 / 7,5 Jahre für die geschlossenen Buckets; der offene "
     ">10Y-Bucket konservativ mit **15 Jahren** (angelehnt an das "
     "Basel-/EBA-Zeitband 10–15 J.; Slotting-Logik analog BCBS 368). "
     "CET1-wirksam ist nur der HfT/FVTPL/FVOCI-Anteil — dieser "
     "IFRS-9-Split ist **je Bank aus der EBA Transparency gemeldet**, "
     "keine Pauschal-Annahme; duration-gewichtet sind über die zehn "
     "Banken ≈ 51 % des Brutto-MtM CET1-wirksam. Konvexität ist "
     "vernachlässigt (linearer Duration-Ansatz — überzeichnet den "
     "Verlust bei großen Zinsanstiegen leicht, wirkt also konservativ). "
     "Der Kanal nutzt den EBA-Perimeter aller zehn Banken und ist "
     "unabhängig davon, ob eine Bank Staaten im IRB-Kreditbuch führt "
     "(Beispiel Santander: Standardansatz im Kreditbuch, vollständig im "
     "Sovereign-Kanal). Durationslogik nach Tuckman/Serrat (2012); die "
     "begleitenden Marktbuch-Analysen des Cockpits — "
     "Staaten-Banken-Verflechtung (Doom Loop) und latente "
     "Bewertungsverluste — stützen sich auf Brunnermeier et al. (2016), "
     "Acharya/Drechsler/Schnabl (2014) und Jiang et al. (2023).")
figure("fig3_bridge.png", 15.8,
       "Abb. 3 · Die CET1-Bridge: Das Kreditbuch erhöht den Nenner (ΔRWA) "
       "und senkt über erwartete Verluste den Zähler (ΔEL); der "
       "Sovereign-Kanal senkt den Zähler zusätzlich (ΔMtM). Entlastende "
       "Gewinn-Effekte werden bewusst nicht angesetzt — die Quote unter "
       "Stress ist eine konservative Untergrenze.")

# ================= 5 · Datenbasis =================
h1("5 · Datenbasis und Integritätssicherung")
h2("Das Banken-Universum")
body("Die zehn größten EU-Banken nach IRB-EAD — zusammen **≈ 8,28 Bio. € "
     "IRB-EAD ≈ 56 %** des gesamten IRB-EAD im EBA-Transparency-Datensatz "
     "(Auswahllogik: Annahme A-08). CET1-Quoten aus dem "
     "EBA-Kapitalpanel, Stichtag 31.12.2024:")
table(
    ["#", "Bank", "Land", "IRB-EAD (Mrd €)", "CET1 31.12.24",
     "Backtest-Zellen"],
    [["1", "Groupe Crédit Agricole ¹", "FR", "1.646", "17,2 %", "27 (+1)*"],
     ["2", "BNP Paribas", "FR", "1.301", "12,9 %", "28"],
     ["3", "ING Groep", "NL", "928", "13,6 %", "24 ²"],
     ["4", "Société Générale", "FR", "791", "13,3 %", "27 (+1)*"],
     ["5", "Groupe BPCE", "FR", "777", "16,2 %", "28"],
     ["6", "Deutsche Bank", "DE", "739", "13,8 %", "28"],
     ["7", "Crédit Mutuel", "FR", "627", "19,4 %", "19 ³"],
     ["8", "Banco Santander", "ES", "595", "12,8 %", "24"],
     ["9", "Coöperatieve Rabobank", "NL", "469", "16,9 %", "24"],
     ["10", "UniCredit", "IT", "409", "16,0 %", "24"]],
    [4, 27, 7, 17, 15, 30], fontsize=8.5)
body("¹ Live-Universum auf Gruppenebene (LEI FR969500TJ5KRTCJQWXH); die "
     "Backtest-Reihe läuft auf der börsennotierten **Crédit Agricole "
     "S.A.** (LEI F0HUI1NY1AZMJMD8LP67, CET1 18,8 %) — der Ebene des "
     "EBA-Transparency-Kapitalpanels. · * (+1) = eine zusätzliche "
     "quellenbelegte, aber modell-ausgeschlossene Audit-Zeile "
     "(`include_in_backtest = 0`, siehe unten). · ² inkl. der "
     "ING-spezifischen Zusatzklasse mortgage_sme; kein QRRE-/Sovereign-"
     "A-IRB im Meldeumfang. · ³ dünnster Meldeumfang (5 Klassen); enthält "
     "die einzige offene Zelle (corporate 2021). Summe genutzter "
     "Backtest-Zeilen: **253**.", size=8.5, italic=True)
h2("Live-Snapshot (Cockpit)")
body("**70 Parameter-Zeilen** = 10 Banken × 7 IRB-Klassen zum einheitlichen "
     "Stichtag **31.12.2024**; 69 der 70 Zeilen tragen PD/LGD-Werte "
     "direkt aus dem EU-CR6-Sub-total des jeweiligen Pillar-3-Reports "
     "mit Seiten- und URL-Beleg (Konfidenzklasse [published]). Die "
     "70. Zeile — Santander-Sovereign — ist ein dokumentierter "
     "Platzhalter ohne IRB-Parameter, weil Santander Staaten "
     "regulatorisch im Standardansatz führt (im Sovereign-MtM-Kanal, "
     "Abschnitt 4, ist Santander dennoch vollständig enthalten).")
h2("Historische Backtest-Reihe")
body("Für die Validierung existiert eine eigene Roh-Reihe: **10 Banken × "
     "FY2021–FY2024, 255 quellenbelegte Datenpunkte** (PD, LGD und EAD je "
     "Klasse aus derselben Report-Vintage). Das entspricht **99,6 % "
     "Abdeckung relativ zum bank-eigenen Meldeumfang** (255 von 256 "
     "gemeldeten Zellen) — Klassen, die eine Bank gar nicht im A-IRB "
     "führt, zählen nicht als Lücke. Davon sind **253 Zeilen "
     "modellfähig**; zwei quellenbelegte Sonderfälle bleiben für Audit "
     "und Abdeckung sichtbar, werden aber wegen Quellen-/Perimeterproblemen "
     "nicht in die Modellrechnung eingespeist (`include_in_backtest = 0`). "
     "Die einzige verbleibende Quellgrenze ist Credit Mutuel Corporate "
     "2021: kein Wert wird aus RWA oder Nachbarjahren abgeleitet.")
h2("Vier Integritäts-Sicherungen je Datenpunkt")
bullet("**Kalibrier-Anker:** Jeder Extraktions-Parser muss zunächst alle "
       "hand-verifizierten 2024-Werte exakt reproduzieren, bevor Vorjahre "
       "gelesen werden.")
bullet("**Dichte-Cross-Check:** RWA ÷ EAD jeder Zeile muss die im Report "
       "ausgewiesene RWA-Dichte treffen.")
bullet("**EAD-Kontinuität:** Sprünge der Exposure-Reihe über die Jahre "
       "werden gegen den Report geprüft (echte Effekte wie IRB→SA-Umstellung "
       "werden dokumentiert, Ausreißer verworfen).")
bullet("**Adversariale Quellprüfung:** Ein unabhängiger Prüfdurchlauf je "
       "Bank gegen die Quell-PDFs (protokolliert in "
       "tools/pillar3_backfill/VERIFICATION_REPORT.json).")

# ================= 6 · Annahmen-Inventar =================
h1("6 · Annahmen-Inventar (A-01 bis A-10)")
body("Jede Annahme trägt eine Konfidenzklasse: **[published]** = "
     "veröffentlichte Messung · **[estimate]** = belegte Schätzung · "
     "**[approximation]** = strukturelle Vereinfachung · **[assumption]** = "
     "gesetzte Konvention. Die Tabelle gibt den Überblick; darunter wird jede Annahme einzeln begründet — inklusive der verworfenen Alternative und der Wirkung einer Fehlspezifikation.")
body("Leselogik: Eine Annahme ist hier nicht automatisch eine frei gesetzte "
     "Zahl. Viele Einträge sind bewusst als [published] klassifiziert, weil "
     "sie direkt aus Offenlegungen oder der CRR stammen. Kritisch sind vor "
     "allem die [estimate]- und [approximation]-Elemente: sie bestimmen, wie "
     "das Modell von beobachtbaren Marktschocks auf nicht direkt "
     "beobachtbare Stress-PD/LGD und Marktwertverluste schließt. Genau diese "
     "Punkte werden deshalb im Cockpit sichtbar gemacht, im Backtest "
     "sensitiviert und in Abschnitt 8 als Scope-Grenzen eingeordnet.")
table(
    ["Prüffeld", "Explizit dokumentierte Annahme / Festlegung",
     "Wo im Dokument?"],
    [["Faktorwahl und Einheiten",
      "Brent als europäisch relevanter Energiepreis-Proxy; Log-Return "
      "ln(P_t/P_0); 10-Jahres-Zins in Prozentpunkten",
      "Abschnitt 2 / A-10"],
     ["Risikoparameter",
      "PD, LGD und EAD aus bankindividuellen Pillar-3-EU-CR6-Zeilen; "
      "Backtest no-look-ahead mit historischen Vintages",
      "Abschnitt 5 / A-01 bis A-03"],
     ["Stress-Transmission",
      "β/γ je Exposure-Klasse, lineare Basiskalibrierung, Floors/Caps und "
      "±50-%-Robustheitstest",
      "Abschnitt 3 / A-04 / Abschnitt 7"],
     ["Regulatorische Kapitalformel",
      "Basel-IRB-ASRF; ρ, 99,9 %-Quantil und Faktor 12,5 regulatorisch "
      "vorgegeben; M = 2,5 Jahre und S = 20 Mio. € explizite Konventionen",
      "Abschnitt 4 / A-05"],
     ["Sovereign-Kanal",
      "Duration-MtM auf Δr, Bucket-Midpoints, >10Y mit 15 Jahren, "
      "CET1-Wirkung nur nach bankindividuellem IFRS-9-Split",
      "Abschnitt 4 / A-06"],
     ["Bewusste Nicht-Modellierung",
      "u. a. Operational Risk, CVA, Spread-Risiko, IFRS-9-Lifetime-EL, "
      "Hedging, Ertrags-Gegeneffekte und bankindividuelle β",
      "Abschnitt 8"]],
    [24, 53, 23], fontsize=8.3)
table(
    ["ID", "Annahme", "Festlegung im Modell", "Klasse", "Quelle/Beleg"],
    [["A-01", "PD-Quelle", "Pillar-3 EU-CR6-Sub-totals je Bank×Klasse, "
      "Stichtag 31.12.2024", "published", "EBA-ITS/2020/04; CRR Art. 431–455"],
     ["A-02", "LGD-Quelle", "identisch A-01 (EU-CR6-LGD)", "published",
      "wie A-01"],
     ["A-02c", "Backtest-Zeitreihe", "eigene Pillar-3-Roh-Reihe, 10 Banken × "
      "FY2021–2024 (255 Quellenzeilen, 253 modellfähig), strikt no-look-ahead", "published",
      "Bank-Pillar-3-Reports FY2021–2024"],
     ["A-03", "EAD", "EBA-Transparency Item 2520522 (post-CCF)", "published",
      "EBA Transparency 2025"],
     ["A-04", "Stress-Transmission", "2-Faktor-β/γ je Klasse (Abschnitt 3); "
      "PD-Floor 3 bp, Cap 50 %; LGD-Floor 5 %", "estimate",
      "EBA 2025 §2.4.2; ECB WP 2897/3112; FSR 5/2024"],
     ["A-05", "Kapitalformel", "Basel-III-IRB (Vasicek/ASRF), α = 99,9 %, "
      "regulatorische Klassenparameter, Laufzeit-Adjustment", "published",
      "BCBS 2017 §272–284; CRR Art. 153/154"],
     ["A-06", "Sovereign-Kanal", "ΔMtM = −Duration·Δr·Exposure; CET1-wirksam "
      "nur FVOCI/FVTPL-Anteil", "approximation",
      "Tuckman/Serrat 2012; EBA-IFRS-9-Split"],
     ["A-07", "Trading Book", "kein eigener Kanal (kleine Handelsbücher, "
      "keine belastbare öffentliche FRTB-Kalibrierung)", "assumption",
      "Scope-Abgrenzung"],
     ["A-08", "Universum", "die 10 größten EU-IRB-Banken (einheitliche "
      "Datenqualität statt Vollabdeckung)", "approximation",
      "EBA Transparency"],
     ["A-09", "Klassen-Raster", "7 IRB-Klassen (corporate, sme_corporate, "
      "mortgage, qrre, other_retail, bank, sovereign)", "approximation",
      "EU-CR6-Struktur"],
     ["A-10", "Faktor-Unabhängigkeit", "ρ(ΔBrent, Δr_10y) = +0,07 → separate "
      "Modellierung ohne Kreuzterm", "estimate",
      "eigene Regression, 1 242 Handelstage"]],
    [7, 17, 40, 12, 24], fontsize=8.5)

h2("Die Annahmen im Detail")

def annahme(aid, titel, text):
    doc.add_paragraph(f"{aid} · {titel}", style="Heading 3")
    body(text, size=9.5, space_after=9)

annahme("A-01", "PD-Quelle: Pillar-3 EU-CR6 (31.12.2024)",
    "**Festlegung:** Je Bank und Klasse die EAD-gewichtete Durchschnitts-PD "
    "aus dem EU-CR6-Sub-total — die regulatorische 1-Jahres-PD nach CRR "
    "Art. 180, inklusive des 100-%-Default-Bands. **Warum:** Es ist die "
    "einzige öffentliche, bankindividuelle und geprüfte PD-Quelle mit "
    "identischer Definition über alle zehn Banken. **Verworfene "
    "Alternative:** PD-Ableitung aus EBA-Transparency-Ausfallquoten "
    "(vermischt Bestands- und Flussgrößen, keine 1-Jahres-PD) sowie "
    "Rating-Agentur-PDs (nicht je Bank × Klasse verfügbar). **Wirkung bei "
    "Fehlspezifikation:** Das PD-Niveau skaliert den erwarteten Verlust "
    "nahezu proportional, die Kapitalanforderung unterproportional — eine "
    "Niveauverzerrung verschiebt also Ergebnisse, kippt aber keine "
    "Richtungsaussage.")
annahme("A-02", "LGD-Quelle: Pillar-3 EU-CR6 (31.12.2024)",
    "**Festlegung:** EAD-gewichtete Durchschnitts-LGD aus derselben "
    "EU-CR6-Tabelle wie A-01. **Warum:** Bankintern geschätzte A-IRB-LGD "
    "spiegeln die tatsächliche Besicherung des Portfolios wider. "
    "**Verworfene Alternative:** die pauschale F-IRB-Aufsichts-LGD von "
    "45 % — sie würde besicherte Portfolien (v. a. Hypotheken) massiv "
    "überzeichnen. **Wirkung bei Fehlspezifikation:** LGD wirkt linear auf "
    "erwarteten Verlust und fast linear auf die Kapitalanforderung — der "
    "direkteste Hebel unter den Parametern.")
annahme("A-02c", "Backtest-Zeitreihe (nur Pillar-3, no-look-ahead)",
    "**Festlegung:** Eigene Roh-Reihe über 10 Banken × FY2021–FY2024 "
    "(255 quellenbelegte Zeilen, davon 253 modellfähig); ein Quartal im "
    "Jahr Y nutzt ausschließlich den Stichtag "
    "31.12.(Y−1) — den jüngsten, der damals publiziert war. **Warum:** Nur "
    "so ist die Validierung echt out-of-sample. **Verworfene Alternative:** "
    "den 2024-Snapshot rückwirkend anzuwenden (Look-ahead-Verzerrung) oder "
    "zwischen Stichtagen zu interpolieren (nicht quellengestützt). "
    "**Wirkung bei Fehlspezifikation:** Mit Look-ahead wäre jedes "
    "Backtest-Ergebnis systematisch geschönt und wertlos.")
annahme("A-03", "EAD: EBA Transparency, post-CCF",
    "**Festlegung:** Exposure at Default je Klasse nach "
    "Kreditumrechnungsfaktoren (post-CCF; EBA-Item 2520522). **Warum:** "
    "Post-CCF ist exakt die Bezugsgröße der IRB-Formel. **Verworfene "
    "Alternative:** Original Exposure — überzeichnet nicht gezogene "
    "Kreditlinien. **Wirkung bei Fehlspezifikation:** EAD wirkt als "
    "Portfoliogewicht linear auf Volumeneffekte, verzerrt aber die "
    "Quoten-Richtung nicht.")
annahme("A-04", "2-Faktor-Transmission (β/γ je Klasse)",
    "**Festlegung:** Lineare Übersetzung der Schocks in ΔPD/ΔLGD über die "
    "Matrix aus Abschnitt 3, mit begründeten Unter-/Obergrenzen. **Warum:** "
    "Die EBA-Stresstest-Methodik sieht sektorale Sensitivitäten explizit "
    "vor; die lineare Form ist transparent, prüfbar und per Regler "
    "übersteuerbar. Die β/γ-Werte sind deshalb keine behaupteten "
    "Punktwahrheiten, sondern eine quellenkonforme Basiskalibrierung, deren "
    "Auswirkung separat robustheitsgeprüft wird. **Verworfene Alternative:** ein Ein-Faktor-Aggregat "
    "(verliert die Sektor-Differenzierung und überzeichnete in Tests die "
    "Magnitude) sowie bankindividuell geschätzte Ökonometrie (Datenlage "
    "öffentlich nicht ausreichend). **Wirkung bei Fehlspezifikation:** Die "
    "β sind der größte Unsicherheitsträger des Modells — deshalb "
    "Konfidenzklasse [estimate] und Sichtbarkeit als eigene Matrix; ein "
    "Fehler skaliert die Stress-Wirkung in etwa proportional.")
annahme("A-05", "IRB-Kapitalformel (Basel III / ASRF)",
    "**Festlegung:** Die unveränderte regulatorische Formel mit α = 99,9 %, "
    "CRR-Klassenparametern und Laufzeit-Adjustment. "
    "Prozessual: Öl-/Zinsschock → PD_stress/LGD_stress → Basel-IRB-K-Formel "
    "→ RWA → CET1. Der Parameter ρ ist dabei kein eigener Modellschritt und "
    "keine von uns geschätzte Annahme, sondern der durch CRR Art. 153/154 "
    "vorgegebene Klassenparameter innerhalb der Formel; er beschreibt die "
    "Kopplung einer Exposure-Klasse an den systematischen Kreditrisikofaktor "
    "des ASRF-Modells, nicht eine empirische Öl-Zins-Korrelation. "
    "**Warum:** Nur der regulatorische Standard macht Modell-Output und "
    "gemeldete RWA vergleichbar. **Verworfene Alternative:** interne "
    "Portfoliomodelle (z. B. Migrationsmatrizen) — nicht gegen die "
    "CET1-Meldung abgleichbar. **Wirkung bei Fehlspezifikation:** Die "
    "Formel selbst ist fixiert; das Modellrisiko liegt vollständig in den "
    "Inputs (A-01 bis A-04).")
annahme("A-06", "Sovereign-Kanal (Duration-MtM, IFRS-9-Split)",
    "**Festlegung:** ΔMtM = −Duration · Δr · Exposure; CET1-wirksam nur der "
    "FVOCI/FVTPL-Anteil je Bank. **Warum:** Das folgt der "
    "IFRS-9-Bewertungslogik — zu fortgeführten Anschaffungskosten (AC) "
    "gehaltene Bestände berühren die CET1-Quote laufend nicht. "
    "**Verworfene Alternative:** volle Marktbewertung aller Bestände "
    "(überzeichnet) oder Verzicht auf den Kanal (unterschlägt das "
    "Zinsrisiko). **Wirkung bei Fehlspezifikation:** Latente AC-Verluste "
    "bleiben unsichtbar — eine bewusste Untererfassung, die das Cockpit "
    "als separate Analyse ausweist (SVB-Lektion, Jiang et al. 2023).")
annahme("A-07", "Kein Trading-Book-Kanal",
    "**Festlegung:** Das Handelsbuch ist kein eigener Stress-Kanal. "
    "**Warum:** Die Handelsbücher der zehn Banken sind relativ klein, und "
    "eine belastbare öffentliche FRTB-Kalibrierung existiert nicht — ein "
    "pauschaler Multiplikator wäre Scheingenauigkeit. **Verworfene "
    "Alternative:** ein Kanal mit angenommenem Marktrisiko-Multiplikator "
    "ohne bankindividuelle Datenbasis. **Wirkung:** Für Banken mit größerem "
    "Handelsbuch wird der Stress tendenziell unterschätzt; diese Grenze ist "
    "als Scope-Grenze dokumentiert (Abschnitt 8).")
annahme("A-08", "Universum: die 10 größten EU-IRB-Banken",
    "**Festlegung:** Genau zehn Institute, ausgewählt nach Größe und "
    "A-IRB-Offenlegungsqualität. **Warum:** Einheitliche, vollständige "
    "Pillar-3-Daten schlagen eine breite, aber heterogene Abdeckung. "
    "**Verworfene Alternative:** alle ~67 IRB-Banken der Transparency "
    "(uneinheitliche Offenlegung, Konzern-/Tochter-Doppelzählungen). "
    "**Wirkung:** Aussagen gelten für diese Gruppe — nicht für den "
    "gesamten EU-Bankenmarkt.")
annahme("A-09", "Klassen-Raster: 7 IRB-Klassen",
    "**Festlegung:** corporate, sme_corporate, mortgage, qrre, "
    "other_retail, bank, sovereign — das EU-CR6-Raster. **Warum:** Es ist "
    "der kleinste gemeinsame Nenner, den alle zehn Banken identisch "
    "offenlegen. **Verworfene Alternative:** feinere Branchenraster "
    "(NACE) — öffentlich nicht flächendeckend verfügbar. **Wirkung:** "
    "Heterogenität innerhalb einer Klasse wird auf den gewichteten "
    "Durchschnitt gemittelt.")
annahme("A-10", "Faktor-Unabhängigkeit (kein Kreuzterm)",
    "**Festlegung:** Öl- und Zinsschock wirken additiv, ohne "
    "Interaktionsterm. **Warum:** Empirisch sind die Faktoren nahezu "
    "unkorreliert (ρ = +0,07 über 1 242 Handelstage). **Verworfene "
    "Alternative:** Copula- oder Interaktionsmodelle — ohne empirischen "
    "Träger nur zusätzliche Komplexität. **Wirkung bei Fehlspezifikation:** "
    "Bei gleichzeitigen Extremschocks entstünde ein Effekt zweiter "
    "Ordnung; im beobachteten Wertebereich ist er vernachlässigbar.")

# ================= 7 · Validierung =================
h1("7 · Validierung: der CET1-Walk-Forward-Backtest")
body("Das Modell wird nicht an dem Datenstand geprüft, mit dem es gebaut "
     "wurde, sondern **rollierend durch die Vergangenheit**: Für jedes Jahr "
     "wird das Portfolio mit den damals bekannten Pillar-3-Parametern des "
     "Vorjahres eingefroren, der tatsächlich eingetretene Öl-/Zins-Schock "
     "des Jahres eingespeist und die prognostizierte CET1-Quote mit der "
     "später gemeldeten verglichen — ohne jeden Blick in die Zukunft.")
figure("fig4_walkforward.png", 16.2,
       "Abb. 4 · No-Look-Ahead-Prinzip: einfrieren (31.12. des Vorjahres) → "
       "realen Schock einspeisen → gegen die gemeldete CET1-Quote halten → "
       "ein Jahr weiterrollen.")
h2("Ergebnis (30 Bank-Jahre, 2022–2024 — vollständiges Panel)")
_sg = STATS["gesamt"]
_sy = STATS["jahre"]
table(
    ["Kennzahl", "Wert", "Bedeutung"],
    [["MAE (mittlerer absoluter Fehler)", f"**{de(_sg['mae'], 2)} pp**",
      "Ø-Abstand Prognose ↔ gemeldete CET1-Quote; auf ~15 %-Niveau "
      "≈ 8 % relativer Abstand"],
     ["Treffer ≤ 1 pp", f"**{de(_sg['within'] * 100, 0)} %**",
      "Anteil der Bank-Jahre mit höchstens 1 Prozentpunkt Abweichung"],
     ["Konservativ-Anteil", f"**{de(_sg['cons'] * 100, 0)} %**",
      "Prognose ≤ gemeldete Quote — das Modell irrt auf der sicheren Seite"],
     ["Bias", f"**{de(_sg['bias'], 2, sign=True)} pp**",
      "systematisch vorsichtig (gewollt: kein "
      "NII-Gegeneffekt angesetzt)"]],
    [30, 14, 56], fontsize=9)
body("Alle 10 Banken sind in allen drei Backtest-Jahren vertreten (30 von "
     "30 möglichen Bank-Jahren). Im Zinsschock-Jahr 2022 liegt das Modell "
     "bei allen 10 Banken auf der konservativen Seite (Ø "
     f"{de(_sy[2022]['bias'], 2, sign=True)} pp — der bewusst nicht "
     "angesetzte NII-Gegeneffekt); in den ruhigeren Jahren 2023/2024 "
     "beträgt der mittlere Fehler "
     f"{de(_sy[2023]['mae'], 2)} bzw. {de(_sy[2024]['mae'], 2)} pp.",
     size=9.5)
h2("Robustheit der Stress-Transmission")
_ss = STATS["sens"]
table(
    ["Test", "MAE", "Bias", "Einordnung"],
    [["β/γ × 0,5", f"{de(_ss[0.5]['mae'], 2)} pp",
      f"{de(_ss[0.5]['bias'], 2, sign=True)} pp",
      "Schwächerer Stress; 2022 bleibt konservativ"],
     ["Basiskalibrierung", f"{de(_sg['mae'], 2)} pp",
      f"{de(_sg['bias'], 2, sign=True)} pp",
      "Hauptfall des Cockpits"],
     ["β/γ × 1,5", f"{de(_ss[1.5]['mae'], 2)} pp",
      f"{de(_ss[1.5]['bias'], 2, sign=True)} pp",
      "Strengerer Stress; bewusst konservative Oberseite"]],
    [22, 14, 14, 50], fontsize=8.5)
body("Der Test skaliert die gesamte PD-/LGD-Transmission um ±50 %. Er "
     "zeigt, ob das Ergebnis nur an einer punktgenauen β-Wahl hängt. Die "
     "Richtung bleibt stabil: stärkere β verschärfen den konservativen "
     "Bias, schwächere β reduzieren ihn, ohne das Backtest-Narrativ zu "
     "drehen.", size=9.5)
h2("Ehrliche Grenze: Point-in-Time vs. Through-the-Cycle")
body("Die einzelnen Melde-Kanäle (PD, Kredit-RWA) sind **nicht "
     "punktprognostizierbar** — und zwar aus einem dokumentierten "
     "Daten-Grund, nicht wegen eines Modellfehlers: Das Modell rechnet "
     "**Point-in-Time** (sofortige Schockreaktion), die gemeldeten "
     "A-IRB-Parameter sind **Through-the-Cycle** — regulatorisch geglättet "
     "und antizyklisch (CRR Art. 180) sowie durch Portfoliosteuerung "
     "überlagert. Beleg 2022: Zins +2,8 pp → Modell-PD +0,5 pp, gemeldete "
     "PD −0,1 pp (ungewichteter Durchschnitt über Bank-Klassen-Paare; "
     "EAD-gewichtet +0,6 vs. −0,2 pp). Validiert ist das Modell deshalb "
     "als **konservatives "
     "Solvenz-/Frühwarn-Instrument** (im CET1-Niveau nah und auf der "
     "sicheren Seite) — nicht als Prognose einzelner Meldegrößen.")

# ================= 8 · Scope =================
h1("8 · Bewusste Scope-Grenzen")
body("Die Grenzen des Modells sind nicht nachträgliche Schwächen, sondern "
     "bewusst gesetzte Abgrenzungen gegen Scheingenauigkeit. Entscheidend "
     "ist jeweils, ob ein Effekt öffentlich, bankindividuell und konsistent "
     "über alle zehn Banken rekonstruierbar ist. Wo das nicht der Fall ist, "
     "wird der Effekt nicht modelliert oder nur als separate Analyse "
     "ausgewiesen.")
table(
    ["Grenze", "Warum nicht modelliert?", "Richtung der Verzerrung"],
    [["Operational Risk",
      "Für einen Öl-/Zins-Einjahresschock liegen keine bankindividuellen "
      "öffentlichen Parameter vor, die eine belastbare dynamische "
      "Operational-Risk-Reaktion erlauben.",
      "Kann Gesamtrisikowirkung unterschätzen."],
     ["CVA / Kontrahentenrisiko",
      "Der Schock ist auf Kreditbuch-PD/LGD und Sovereign-MtM kalibriert; "
      "Derivate- und Gegenparteinetting sind aus Offenlegungen nicht "
      "durchgängig rekonstruierbar.",
      "Kann Kapitalbelastung in Marktstressphasen unterschätzen."],
     ["Sovereign-Spread-Risiko",
      "Modelliert wird der Niveauzins; länderspezifische Spreadpfade "
      "(z. B. Italien vs. Bund) wären eine zusätzliche Szenarioannahme.",
      "Unterschätzt Stress bei peripheren Staatsanleihen."],
     ["IFRS-9-Lifetime-EL / Stage-2",
      "Das Modell ist ein Einjahres-Solvency-Stress. Eine Lifetime-EL-"
      "Simulation benötigt migrations- und portfoliointerne Daten.",
      "Konservativer ΔEL-Proxy, aber keine vollständige IFRS-9-Prognose."],
     ["Mehrjahres-Stresspfade",
      "Der Backtest und das Cockpit sind auf ein Jahr ausgelegt; "
      "3-Jahres-EBA-Pfade würden Bilanzmanagement, Neugeschäft und "
      "Kapitalmaßnahmen erfordern.",
      "Keine Aussage zur Pfaddynamik nach Jahr 1."],
     ["Hedging",
      "Swaps, Futures, CDS und Hedge Accounting sind öffentlich nicht "
      "hinreichend granular je Bank und Exposure rekonstruierbar.",
      "Kann Verluste überschätzen, wenn Banken stark gehedgt sind."],
     ["Ertrags-Gegeneffekte",
      "NII, Gebühren und Managementmaßnahmen werden nicht gegengerechnet, "
      "weil sie keine reine Risikoparameter-Transmission sind.",
      "CET1-Stress ist bewusst konservativ; 2022 zeigt diesen Effekt."],
     ["Bank-individuelle β",
      "Alle Banken teilen je Exposure-Klasse dieselbe Sensitivität; "
      "Bankunterschiede entstehen über echte PD/LGD/EAD-Mixe.",
      "Kann idiosynkratische Geschäftsmodelle glätten."]],
    [22, 54, 24], fontsize=8.3)
body("Kurz gesagt: Das Modell ist ein **konservatives, datenprüfbares "
     "Solvenz-Stressmodell**. Es ist nicht als vollständige Bankplanung, "
     "IFRS-9-Forecast, Handelsbuchmodell oder Anlageempfehlung zu lesen. "
     "Seine Stärke liegt in der transparenten, reproduzierbaren Kette von "
     "offengelegten Risikoparametern über regulatorische Kapitalformeln bis "
     "zur CET1-Quote.", italic=True, size=9.5)

# ================= 9 · Bibliographie =================
h1("9 · Bibliographie")
body("Vollständige Liste aller im Modell verwendeten Quellen — gegliedert "
     "nach Regulatorik, wissenschaftlicher Literatur und Datenquellen. Jede "
     "Quelle ist an der Stelle ihrer Verwendung im Dokument bzw. im "
     "Modell-Code referenziert.")
h2("Regulatorische Quellen")
for b in [
    "BCBS (2017). Basel III: Finalising post-crisis reforms. Basel Committee "
    "on Banking Supervision, §272–284 (IRB-Kapitalformel).",
    "Board of Governors / OCC (2011). SR 11-7: Supervisory Guidance on Model "
    "Risk Management (Outcomes-Analysis des Backtests).",
    "EBA (2020). ITS/2020/04 — Implementing Technical Standards on public "
    "disclosures (EU-CR6-Template, Pillar-3-Datenbasis).",
    "EBA (2024). 2025 EU-wide Stress Test — Methodological Note, §2.4.2 "
    "(sektorale Sensitivitäten als aufsichtlicher Projektionsansatz).",
    "EBA/ESRB (2025). 2025 EU-wide Stress Test — Macro-financial Scenario, "
    "§4.1.6 (adverse Zinspfade als Plausibilitätsanker).",
    "EBA (2025). 2025 EU-wide Stress Test — Results, Fig. 22 (relative "
    "Verlustquoten je Portfolio als Quervalidierung der β-Struktur).",
    "EBA (2025). Report on the 2024 Credit Risk Benchmarking Exercise "
    "(PD-/LGD-Niveaus je IRB-Klasse).",
    "EBA GL 2014/14. Guidelines on ICAAP/Stress-Testing (Governance-Rahmen).",
    "Verordnung (EU) Nr. 575/2013 (CRR) — insb. Art. 153/154 "
    "(IRB-Risikogewichte), Art. 180/181 (PD-/LGD-Schätzung, TTC-Charakter), "
    "Art. 431–455 (Offenlegung).",
]:
    bullet(b)
h2("Wissenschaftliche Literatur")
for b in [
    "Acharya, V., Drechsler, I. & Schnabl, P. (2014). A Pyrrhic Victory? "
    "Bank Bailouts and Sovereign Credit Risk. Journal of Finance "
    "(Staaten-Banken-Verflechtung, Marktbuch-Analyse).",
    "Bandoni, E., Fourné, M. & Jarmulska, B. (2025). Mortgage loan rates and "
    "the defaults of variable rate mortgages. ECB Working Paper 3112 "
    "(Zins-Sensitivität der Hypotheken-PD).",
    "Brunnermeier, M. et al. (2016). The Sovereign-Bank Diabolic Loop and "
    "ESBies. American Economic Review P&P (Doom-Loop, Marktbuch-Analyse).",
    "Gordy, M. (2003). A Risk-Factor Model Foundation for Ratings-Based Bank "
    "Capital Rules. Journal of Financial Intermediation (ASRF-Fundierung).",
    "Hyndman, R. J. & Athanasopoulos, G. (2021). Forecasting: Principles and "
    "Practice, 3. Aufl., Kap. 5.8 (Prognose-Evaluation, MAE).",
    "Jiang, E., Matvos, G., Piskorski, T. & Seru, A. (2023). Monetary "
    "Tightening and U.S. Bank Fragility. NBER Working Paper 31048 (latente "
    "Bewertungsverluste, Marktbuch-Analyse).",
    "Konietschke, P., Metzler, J. & Ponte Marques, A. (2026). A quantile "
    "probability model for sectoral corporate defaults in Europe. ECB "
    "Working Paper 3207 (Sektor-Heterogenität der Unternehmens-PD).",
    "Lo Duca, M., Moccero, D. & Parlapiano, F. (2024). The impact of "
    "macroeconomic and monetary policy shocks on credit risk in the euro "
    "area corporate sector. ECB Working Paper 2897 (Corporate-/KMU-β).",
    "Nelson, C. & Siegel, A. (1987). Parsimonious Modeling of Yield Curves. "
    "Journal of Business (Zinsstrukturmodell).",
    "Pesaran, M. H. & Timmermann, A. (1992). A Simple Nonparametric Test of "
    "Predictive Performance. Journal of Business & Economic Statistics "
    "(Richtungs-Trefferquote).",
    "Svensson, L. (1994). Estimating and Interpreting Forward Interest "
    "Rates: Sweden 1992–1994. NBER Working Paper 4871 "
    "(Zinsstrukturmodell der Bundesbank-Parameter).",
    "Tuckman, B. & Serrat, A. (2012). Fixed Income Securities, 3. Aufl. "
    "(Modified Duration, Sovereign-Kanal).",
    "Vasicek, O. (1991). Limiting Loan Loss Probability Distribution. KMV "
    "Working Paper (Portfolio-Verlustverteilung).",
    "Vasicek, O. (2002). Loan Portfolio Value. Risk Magazine, Dezember "
    "(geschlossene Form der IRB-Formel).",
    "ECB (2024). Financial Stability Review, Mai 2024 (Energie-/Lebens"
    "haltungskosten-Wirkung auf Haushalte; Retail-β).",
]:
    bullet(b)
h2("Datenquellen")
for b in [
    "Pillar-3-Berichte (EU CR6) der zehn Banken, FY2021–FY2024 — "
    "Risikoparameter PD/LGD/EAD (je Zeile mit Seiten-/URL-Beleg).",
    "EBA EU-wide Transparency Exercise (Vintages 2020–2025) — realisierte "
    "CET1- und RWA-Werte, IFRS-9-Split der Staatsanleihen.",
    "ICE Brent Crude (täglich) — Ölpreis-Faktor.",
    "Deutsche Bundesbank — täglich geschätzte Svensson-Parameter der "
    "Zinsstrukturkurve (10-Jahres-Zins-Faktor).",
]:
    bullet(b)

# ---- Word-Einstellungen: bestFit-Zoom. Das Inhaltsverzeichnis wird im
#      Build-/QA-Schritt per Word-COM aktualisiert, nicht automatisch beim
#      Öffnen. Das vermeidet blockierende Update-Dialoge in fremden
#      Word-Sessions.
_settings = doc.settings.element
_zoom = _settings.find(qn("w:zoom"))
if _zoom is None:
    _zoom = OxmlElement("w:zoom")
    _settings.insert(0, _zoom)
_zoom.set(qn("w:val"), "bestFit")
_zoom.set(qn("w:percent"), "100")

doc.save(OUT)
print("geschrieben:", OUT, "| Live-Stats:", STATS.get("live"))
